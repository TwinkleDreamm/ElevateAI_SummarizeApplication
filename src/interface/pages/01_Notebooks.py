from __future__ import annotations

import streamlit as st
from typing import List, Dict, Any
from datetime import datetime
import threading
import time
import re

from src.interface.notebooks import store
from src.interface.app_context import get_context
from src.interface.notebooks.ingest import ingest_uploaded_files, ingest_url
from src.interface.utils.notebook_helper import NotebookHelper
import hashlib
import json
import os
from pathlib import Path
from src.utils.logger import logger

# Thread-safe in-memory task status (avoid using Streamlit APIs inside threads)
_TASK_STATUS: dict[str, dict] = {}

def _sanitize_for_filename(name: str) -> str:
    """Make a safe filename component from notebook name (ASCII-ish, underscores)."""
    try:
        import re
        safe = re.sub(r"\s+", "_", name.strip()) if name else "Notebook"
        safe = re.sub(r"[^\w\-]+", "", safe)
        return safe or "Notebook"
    except Exception:
        return "Notebook"

def _strip_markdown_for_docx(text: str) -> str:
    """Remove common Markdown decorations for DOCX plain text output."""
    try:
        import re
        if not text:
            return ""
        t = text
        # Inline code/backticks
        t = t.replace("```", "").replace("`", "")
        # Bold/italic markers
        t = t.replace("**", "").replace("__", "").replace("*", "")
        # Headings at line start like #, ##, ###
        t = re.sub(r"(?m)^\s*#{1,6}\s*", "", t)
        # Convert markdown links [text](url) -> text (url)
        t = re.sub(r"\[([^\]]+)\]\(([^\)]+)\)", r"\1 (\2)", t)
        # Remove image syntax ![alt](url) -> alt (url)
        t = re.sub(r"!\[([^\]]*)\]\(([^\)]+)\)", r"\1 (\2)", t)
        # Normalize bullet prefixes to a simple dash
        t = re.sub(r"(?m)^\s*[-*•]\s+", "- ", t)
        return t.strip()
    except Exception:
        return text or ""

def _get_notes_storage_path(notebook_id: str) -> Path:
    """Get the storage path for notes of a specific notebook."""
    data_dir = Path("data/notes")
    data_dir.mkdir(parents=True, exist_ok=True)
    return data_dir / f"notebook_{notebook_id}_notes.json"


def _save_notes_to_storage(notebook_id: str, notes: list):
    """Save notes to persistent storage."""
    try:
        storage_path = _get_notes_storage_path(notebook_id)
        with open(storage_path, 'w', encoding='utf-8') as f:
            json.dump(notes, f, ensure_ascii=False, indent=2, default=str)
    except Exception as e:
        st.error(f"Failed to save notes to storage: {e}")


def _load_notes_from_storage(notebook_id: str) -> list:
    """Load notes from persistent storage."""
    try:
        storage_path = _get_notes_storage_path(notebook_id)
        if storage_path.exists():
            with open(storage_path, 'r', encoding='utf-8') as f:
                return json.load(f)
    except Exception as e:
        st.error(f"Failed to load notes from storage: {e}")
    return []


def _get_notebook_folder(notebook_id: str) -> Path:
    """Ensure and return the per-notebook working folder under data/notebooks/{id}."""
    root = Path("data/notebooks") / str(notebook_id)
    root.mkdir(parents=True, exist_ok=True)
    return root


def _get_latest_docx_path(notebook_id: str) -> Path:
    """Return latest DOCX report path if exists (Report_*.docx), fallback to legacy."""
    folder = _get_notebook_folder(notebook_id)
    latest = _find_latest_file(folder, ["Report_*.docx"]) or folder / "overview_latest.docx"
    return latest


def _get_latest_audio_path(notebook_id: str) -> Path:
    """Return latest MP3 report path if exists (Report_*.mp3), fallback to legacy."""
    folder = _get_notebook_folder(notebook_id)
    latest = _find_latest_file(folder, ["Report_*.mp3"]) or folder / "overview_latest.mp3"
    return latest


def _purge_old_overview_files(notebook_id: str, patterns: list[str], keep: list[Path]) -> None:
    """Delete old overview files matching patterns except those in keep list."""
    try:
        folder = _get_notebook_folder(notebook_id)
        keep_set = {p.resolve() for p in keep}
        for pat in patterns:
            for fp in folder.glob(pat):
                try:
                    if fp.resolve() not in keep_set:
                        fp.unlink()
                        logger.info(f"[StudioOverview] Purged old file: {fp.name}")
                except Exception:
                    continue
    except Exception:
        pass


def _get_latest_mindmap_path(notebook_id: str) -> Path:
    """Return stable path placeholder for future mindmap exports."""
    return _get_notebook_folder(notebook_id) / "mindmap_latest.png"

def _collect_notebook_texts(notebook_id: str, *, vector_db=None, search_engine=None) -> str:
    """Collect and join all texts for a notebook from the metadata store."""
    try:
        if vector_db is None:
            ctx = get_context()
            vector_db = ctx.get("vector_db")
        if not vector_db:
            logger.warning(f"[StudioOverview] No vector_db available for notebook {notebook_id}")
            return ""
        metas = vector_db.metadata_manager.search_metadata({"notebook_id": notebook_id})
        logger.info(f"[StudioOverview] Collected {len(metas)} metadata entries for notebook {notebook_id}")
        texts = [m.get("text", "") for m in metas if m.get("text")]
        # Deduplicate and join with separators
        seen = set()
        unique_texts = []
        for t in texts:
            if t not in seen:
                unique_texts.append(t)
                seen.add(t)
        combined = "\n\n---\n\n".join(unique_texts)
        if combined.strip():
            logger.info(f"[StudioOverview] Combined text length: {len(combined)} chars (from metadata)")
            return combined

        # Fallback 1: load saved Studio notes from storage
        try:
            notes_path = _get_notes_storage_path(notebook_id)
            if notes_path.exists():
                notes = json.loads(notes_path.read_text(encoding="utf-8"))
                note_texts = [n.get("content", "") for n in notes if n.get("content")]
                if note_texts:
                    logger.info(f"[StudioOverview] Using {len(note_texts)} Studio notes as content source")
                    return "\n\n---\n\n".join(note_texts)
        except Exception:
            pass

        # Fallback 2: use cached overview from store
        try:
            ov = store.get_overview(notebook_id)
            if ov and ov.strip():
                logger.info("[StudioOverview] Using cached overview from store as content source")
                return ov
        except Exception:
            pass

        # Fallback: pull content via semantic search using broad queries
        if search_engine is None and vector_db is not None:
            try:
                from src.search.semantic_search import SemanticSearchEngine
                search_engine = SemanticSearchEngine(vector_db=vector_db)
            except Exception:
                search_engine = None
        if search_engine is None:
            logger.warning(f"[StudioOverview] No search engine available for notebook {notebook_id}")
            return ""
        queries = [
            "tóm tắt nội dung",
            "nội dung chính",
            "overview",
            "key points",
        ]
        collected = []
        for q in queries:
            try:
                results = search_engine.search(q, k=50, threshold=-1.0, filters={"notebook_id": notebook_id})
                logger.info(f"[StudioOverview] Search '{q}' returned {len(results)} results")
                for r in results:
                    txt = r.text or r.metadata.get("text", "")
                    if txt and txt not in collected:
                        collected.append(txt)
            except Exception:
                pass
        logger.info(f"[StudioOverview] Collected {len(collected)} snippets via search; combined length {sum(len(x) for x in collected)}")
        return "\n\n---\n\n".join(collected)
    except Exception:
        logger.error(f"[StudioOverview] Failed collecting texts for notebook {notebook_id}")
        return ""


def _find_latest_file(folder: Path, patterns: list[str]) -> Path | None:
    """Find the most recently modified file in folder matching any of patterns."""
    try:
        candidates: list[Path] = []
        for p in patterns:
            candidates.extend(sorted(folder.glob(p)))
        if not candidates:
            return None
        return max(candidates, key=lambda p: p.stat().st_mtime)
    except Exception:
        return None


def _run_langgraph_summary(
    full_text: str,
    *,
    prompt_manager=None,
    llm_client=None,
    additional_instructions: str | None = None,
    max_tokens: int | None = None,
) -> str:
    """Run LangGraph summarization workflow (with safe fallback) and return summary text."""
    if not full_text:
        return ""
    try:
        from src.ai.langgraph.workflows.summarization_workflow import create_summarization_workflow
    except Exception:
        create_summarization_workflow = None  # type: ignore
    try:
        if prompt_manager is None or llm_client is None:
            ctx = get_context()
            prompt_manager = ctx.get("prompt_manager")
            llm_client = ctx.get("llm_client")
        if not prompt_manager or not llm_client:
            return full_text[:50000]
        workflow = create_summarization_workflow(recursion_limit=8) if create_summarization_workflow else None
        # Determine instruction profile
        detailed_default = (
            "Viết bản tổng quan CHI TIẾT, mạch lạc bằng tiếng Việt (≈800–1500 từ). "
            "Giữ cấu trúc một báo cáo hoàn chỉnh: (1) Mở đầu ngắn gọn (1–3 câu) nêu mục tiêu/nội dung chính; "
            "(2) Các mục nội dung theo chủ đề (dùng tiêu đề ngắn), dưới mỗi mục liệt kê bullets nêu ý chính, số liệu quan trọng, ví dụ minh họa; "
            "(3) Kết luận và khuyến nghị/next steps. Không tách theo từng nguồn; hợp nhất thông tin từ mọi nguồn thành một bức tranh thống nhất. "
            "Khi có thể, chèn trích dẫn ngắn gọn trong ngoặc vuông [Nguồn: tiêu đề/URL hoặc mã ghi chú] ngay sau dữ kiện/số liệu. "
            "Văn phong rõ ràng, súc tích, chuyên nghiệp, ưu tiên tính chính xác."
        )
        instructions = additional_instructions or detailed_default
        is_concise = "150–250" in instructions or "150–250" in instructions
        gen_max_tokens = max_tokens if isinstance(max_tokens, int) and max_tokens > 0 else (600 if is_concise else 1600)
        if workflow is None:
            # Fallback: single-pass summary using prompt manager + llm_client
            messages = prompt_manager.build_summary_prompt(
                content=full_text[:50000],
                additional_instructions=instructions,
            )
            resp = llm_client.generate_response(messages, max_tokens=gen_max_tokens, temperature=0.2)
            content = resp.get("content", "") if isinstance(resp, dict) else getattr(resp, "content", "")
            return content
        state = workflow.invoke({
            "content": full_text,
            "llm_client": llm_client,
            "prompt_manager": prompt_manager,
            "remaining_loops": 2,
            "additional_instructions": instructions,
            "max_tokens": gen_max_tokens,
        })
        summary = state.get("summary", "")
        if not summary or not summary.strip():
            # Secondary fallback with stricter prompt
            messages = prompt_manager.build_summary_prompt(
                content=full_text[:50000],
                additional_instructions=instructions,
            )
            resp = llm_client.generate_response(messages, max_tokens=gen_max_tokens, temperature=0.2)
            return resp.get("content", "") if isinstance(resp, dict) else getattr(resp, "content", "")
        # Ensure final overview length aligns with instruction profile
        try:
            def _word_count(t: str) -> int:
                return len((t or "").split())
            if ((not is_concise and _word_count(summary) > 1100) or (is_concise and _word_count(summary) > 330)) and prompt_manager and llm_client:
                instruction = (
                    "Rút gọn nội dung sau thành "
                    + ("tổng quan khoảng 150–250 từ" if is_concise else "báo cáo khoảng 800–1500 từ")
                    + ", giữ cấu trúc mạch lạc: mở đầu ngắn; các mục theo chủ đề với bullets nêu ý chính, số liệu, ví dụ; "
                    "kết luận/khuyến nghị. Không tách theo nguồn; hợp nhất thông tin; giữ trích dẫn dạng [Nguồn: ...] nếu có."
                )
                messages = (
                    prompt_manager.build_generic_prompt(
                        system_instruction=instruction,
                        user_content=summary[:50000],
                    ) if hasattr(prompt_manager, "build_generic_prompt") else [
                        {"role": "system", "content": instruction},
                        {"role": "user", "content": summary[:50000]},
                    ]
                )
                resp = llm_client.generate_response(messages, max_tokens=gen_max_tokens, temperature=0.2)
                condensed = resp.get("content", "") if isinstance(resp, dict) else getattr(resp, "content", "")
                if condensed:
                    return condensed
        except Exception:
            pass
        return summary
    except Exception:
        return full_text[:50000]


def _background_generate_docx_overview(notebook_id: str, task_key: str) -> None:
    """Worker: generate DOCX overview and update session state when done."""
    try:
        # Build isolated context (avoid Streamlit session access in thread)
        from src.interface.app_context import _build_context as _build_isolated_ctx  # type: ignore
        ctx = _build_isolated_ctx()
        logger.info(f"[StudioOverview] (DOCX) Thread start for notebook {notebook_id}")
        # Gather content and summarize
        full_text = _collect_notebook_texts(
            notebook_id,
            vector_db=ctx.get("vector_db"),
            search_engine=ctx.get("search_engine"),
        )
        logger.info(f"[StudioOverview] (DOCX) Collected text length: {len(full_text)}")
        summary = _run_langgraph_summary(
            full_text,
            prompt_manager=ctx.get("prompt_manager"),
            llm_client=ctx.get("llm_client"),
            additional_instructions=(
                "Xuất bản báo cáo DOCX: Trình bày văn bản mạch lạc, không dùng Markdown (không **, __, #, ```). "
                "Sử dụng đoạn văn và dấu gạch đầu dòng đơn giản khi cần; tránh ký hiệu đặc biệt. "
                "Nội dung chi tiết, rõ ràng, chuyên nghiệp theo cấu trúc báo cáo."
            ),
        )
        if not summary or not summary.strip():
            logger.warning("[StudioOverview] (DOCX) Empty summary from LLM; falling back to raw text slice")
            summary = (full_text or "")

        # Create DOCX with timestamped Report_{name}_{datetime}.docx and update latest alias for UI
        nb_folder = _get_notebook_folder(notebook_id)
        try:
            nb = store.get_notebook(notebook_id)
            nb_name = getattr(nb, "name", "Notebook") if nb else "Notebook"
        except Exception:
            nb_name = "Notebook"
        safe_name = _sanitize_for_filename(nb_name)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_path: Path = nb_folder / f"Report_{safe_name}_{timestamp}.docx"
        try:
            from docx import Document  # type: ignore
            doc = Document()
            title = "Notebook Overview"
            try:
                nb = store.get_notebook(notebook_id)
                if nb and getattr(nb, "name", None):
                    title = f"Notebook Overview — {nb.name}"
            except Exception:
                pass
            doc.add_heading(title, level=1)
            if summary:
                doc.add_paragraph(_strip_markdown_for_docx(summary))
            else:
                doc.add_paragraph("No content available to summarize.")
            doc.save(str(file_path))
        except Exception:
            # Fallback to plain text file if python-docx is unavailable
            file_path = nb_folder / f"Report_{safe_name}_{timestamp}.txt"
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(_strip_markdown_for_docx(summary) if summary else "No content available to summarize.")

        # Purge legacy timestamped files by old pattern; keep new file
        _purge_old_overview_files(notebook_id, ["overview_*.docx", "overview_*.txt"], [file_path])

        # Update status (in-memory dict only)
        _TASK_STATUS[task_key] = {"running": False, "file_path": str(file_path), "error": None}
        logger.info(f"[StudioOverview] (DOCX) Written overview to {file_path}")
    except Exception as e:
        _TASK_STATUS[task_key] = {"running": False, "file_path": None, "error": str(e)}
        logger.error(f"[StudioOverview] (DOCX) Failed: {e}")


def _background_generate_audio_overview(notebook_id: str, task_key: str) -> None:
    """Worker: generate audio overview (MP3) and update session state when done."""
    try:
        # Build isolated context (avoid Streamlit session access in thread)
        from src.interface.app_context import _build_context as _build_isolated_ctx  # type: ignore
        ctx = _build_isolated_ctx()
        logger.info(f"[StudioOverview] (AUDIO) Thread start for notebook {notebook_id}")
        # Gather content and summarize
        full_text = _collect_notebook_texts(
            notebook_id,
            vector_db=ctx.get("vector_db"),
            search_engine=ctx.get("search_engine"),
        )
        logger.info(f"[StudioOverview] (AUDIO) Collected text length: {len(full_text)}")
        summary = _run_langgraph_summary(
            full_text,
            prompt_manager=ctx.get("prompt_manager"),
            llm_client=ctx.get("llm_client"),
        )
        if not summary or not summary.strip():
            logger.warning("[StudioOverview] (AUDIO) Empty summary from LLM; falling back to raw text slice")
            summary = (full_text or "")[:4000]

        tts_client = ctx.get("tts_client")
        voice = "alloy"
        model = "tts-1"
        if not tts_client:
            raise RuntimeError("TTS service not available")
        audio_bytes = tts_client.text_to_speech(
            text=summary[:4000] if summary else "No content available to summarize.",
            voice=voice,
            model=model,
            instructions="Speak clearly and at a moderate pace, suitable for an overview report."
        )
        if not audio_bytes:
            raise RuntimeError("Failed to synthesize audio")

        nb_folder = _get_notebook_folder(notebook_id)
        try:
            nb = store.get_notebook(notebook_id)
            nb_name = getattr(nb, "name", "Notebook") if nb else "Notebook"
        except Exception:
            nb_name = "Notebook"
        safe_name = _sanitize_for_filename(nb_name)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_path = nb_folder / f"Report_{safe_name}_{timestamp}.mp3"
        with open(file_path, "wb") as f:
            f.write(audio_bytes)

        # Purge legacy timestamped files by old patterns; keep new file
        _purge_old_overview_files(notebook_id, ["overview_*.mp3", "overview_*.wav", "overview_*.m4a"], [file_path])

        _TASK_STATUS[task_key] = {"running": False, "file_path": str(file_path), "error": None}
        logger.info(f"[StudioOverview] (AUDIO) Written audio to {file_path}")
    except Exception as e:
        _TASK_STATUS[task_key] = {"running": False, "file_path": None, "error": str(e)}
        logger.error(f"[StudioOverview] (AUDIO) Failed: {e}")


def _generate_example_questions_from_summary(summary: str, *, prompt_manager=None, llm_client=None, target_lang: str = "vi") -> list[str]:
    """Generate 3 example questions using LLM based on the provided summary.

    Returns a list of up to 3 questions, short and diverse. Fallback to defaults if LLM fails.
    """
    try:
        if not prompt_manager or not llm_client or not summary:
            raise RuntimeError("Missing components or empty summary")
        if target_lang.startswith("vi"):
            instruction = (
                "Tạo đúng 3 câu hỏi ví dụ ngắn gọn, thông minh, đa dạng kiểu (tóm tắt, phân tích, so sánh) "
                "dựa trên phần tóm tắt dưới đây. Trả về mỗi câu trên một dòng, không đánh số, tiếng Việt."
            )
        else:
            instruction = (
                "Generate exactly 3 short, smart, diverse example questions (summary, analysis, comparison) "
                "based on the summary below. Return one question per line, no numbering, in English."
            )
        messages = prompt_manager.build_generic_prompt(
            system_instruction=instruction,
            user_content=summary[:6000],
        ) if hasattr(prompt_manager, "build_generic_prompt") else [
            {"role": "system", "content": instruction},
            {"role": "user", "content": summary[:6000]},
        ]
        resp = llm_client.generate_response(messages, max_tokens=256, temperature=0.3)
        content = resp.get("content", "") if isinstance(resp, dict) else getattr(resp, "content", "")
        lines = [ln.strip() for ln in content.splitlines() if ln.strip()]
        # Filter out numbered bullets and take first 3
        cleaned = []
        for ln in lines:
            if ln[:2] in {"1.", "2.", "3."} or ln[:1] in {"-", "*", "•"}:
                ln = ln[2:].strip() if ln[:2] in {"1.", "2.", "3."} else ln[1:].strip()
            if ln:
                cleaned.append(ln)
            if len(cleaned) >= 3:
                break
        if len(cleaned) >= 3:
            return cleaned[:3]
    except Exception:
        pass
    # Fallback
    return [
        "Tóm tắt 5 ý chính của tài liệu này, kèm trích dẫn nguồn.",
        "Liệt kê mốc thời gian quan trọng và nguồn trích dẫn.",
        "So sánh hai quan điểm chính trong các nguồn, kèm trích dẫn.",
    ]


def _background_generate_overview_and_examples(notebook_id: str, overview_task_key: str, examples_task_key: str) -> None:
    """Worker: generate text overview and example questions using LangChain without blocking UI."""
    try:
        from src.interface.app_context import _build_context as _build_isolated_ctx  # type: ignore
        ctx = _build_isolated_ctx()
        logger.info(f"[StudioOverview] (TEXT) Thread start for notebook {notebook_id}")
        full_text = _collect_notebook_texts(
            notebook_id,
            vector_db=ctx.get("vector_db"),
            search_engine=ctx.get("search_engine"),
        )
        logger.info(f"[StudioOverview] (TEXT) Collected text length: {len(full_text)}")

        # Generate overview via LangGraph/LLM
        summary = _run_langgraph_summary(
            full_text,
            prompt_manager=ctx.get("prompt_manager"),
            llm_client=ctx.get("llm_client"),
            additional_instructions=(
                "Viết một bản tổng quan ngắn gọn bằng tiếng Việt, độ dài khoảng 150–250 từ. "
                "Cấu trúc: 1–2 câu mở đầu rất ngắn; sau đó các gạch đầu dòng nêu 4–6 ý chính; "
                "có thể thêm 1 đoạn ngắn kết luận nếu cần. Ngắn gọn, rõ ràng, tránh chi tiết thừa."
            ),
            max_tokens=600,
        )
        if not summary or not summary.strip():
            nb = None
            try:
                nb = store.get_notebook(notebook_id)
            except Exception:
                pass
            summary = (getattr(nb, "description", None) or full_text or "").strip()[:4000]

        # Persist overview
        try:
            store.update_overview(notebook_id, summary)
        except Exception:
            pass
        _TASK_STATUS[overview_task_key] = {"running": False, "result": summary, "error": None}

        # Detect language from content to guide examples
        try:
            vi_chars = sum(ch in "àáạảãâầấậẩẫăằắặẳẵèéẹẻẽêềếệểễìíịỉĩòóọỏõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹđ" for ch in (full_text or "").lower())
            target_lang = "vi" if vi_chars > 0 else "en"
        except Exception:
            target_lang = "vi"

        # Generate examples via LLM
        questions = _generate_example_questions_from_summary(
            summary,
            prompt_manager=ctx.get("prompt_manager"),
            llm_client=ctx.get("llm_client"),
            target_lang=target_lang,
        )
        try:
            store.update_examples(notebook_id, questions)
        except Exception:
            pass
        _TASK_STATUS[examples_task_key] = {"running": False, "result": questions, "error": None}
        logger.info(f"[StudioOverview] (TEXT) Overview and examples generated for notebook {notebook_id}")
    except Exception as e:
        _TASK_STATUS[overview_task_key] = {"running": False, "result": None, "error": str(e)}
        _TASK_STATUS[examples_task_key] = {"running": False, "result": None, "error": str(e)}
        logger.error(f"[StudioOverview] (TEXT) Failed: {e}")

def _cleanup_old_overview_files(notebook_id: str, keep_latest: int = 1) -> dict[str, int]:
    """Clean up old overview files, keeping only the latest ones per type."""
    try:
        nb_folder = _get_notebook_folder(notebook_id)
        if not nb_folder.exists():
            return {"deleted": 0, "kept": 0}
        
        # Group files by type and find the latest ones
        file_groups = {
            "docx": [],
            "txt": [],
            "mp3": [],
            "wav": [],
            "m4a": []
        }
        
        for file_path in nb_folder.iterdir():
            if file_path.is_file():
                suffix = file_path.suffix.lower()
                if suffix in file_groups:
                    file_groups[suffix].append(file_path)
        
        deleted_count = 0
        kept_count = 0
        
        for file_type, files in file_groups.items():
            if len(files) <= keep_latest:
                kept_count += len(files)
                continue
                
            # Sort by modification time (newest first)
            sorted_files = sorted(files, key=lambda p: p.stat().st_mtime, reverse=True)
            
            # Keep the latest files
            files_to_keep = sorted_files[:keep_latest]
            files_to_delete = sorted_files[keep_latest:]
            
            # Delete old files
            for old_file in files_to_delete:
                try:
                    old_file.unlink()
                    deleted_count += 1
                    logger.info(f"[Cleanup] Deleted old {file_type} file: {old_file.name}")
                except Exception as e:
                    logger.warning(f"[Cleanup] Failed to delete {old_file.name}: {e}")
            
            kept_count += len(files_to_keep)
        
        logger.info(f"[Cleanup] Notebook {notebook_id}: deleted {deleted_count} old files, kept {kept_count} latest files")
        return {"deleted": deleted_count, "kept": kept_count}
        
    except Exception as e:
        logger.error(f"[Cleanup] Failed to cleanup notebook {notebook_id}: {e}")
        return {"deleted": 0, "kept": 0, "error": str(e)}


def _get_storage_info(notebook_id: str) -> dict[str, Any]:
    """Get storage information for a notebook's overview files."""
    try:
        nb_folder = _get_notebook_folder(notebook_id)
        if not nb_folder.exists():
            return {"total_size": 0, "file_count": 0, "file_types": {}}
        
        total_size = 0
        file_count = 0
        file_types = {}
        
        for file_path in nb_folder.iterdir():
            if file_path.is_file():
                try:
                    file_size = file_path.stat().st_size
                    total_size += file_size
                    file_count += 1
                    
                    suffix = file_path.suffix.lower()
                    if suffix not in file_types:
                        file_types[suffix] = {"count": 0, "size": 0}
                    file_types[suffix]["count"] += 1
                    file_types[suffix]["size"] += file_size
                except Exception:
                    pass
        
        return {
            "total_size": total_size,
            "file_count": file_count,
            "file_types": file_types,
            "folder_path": str(nb_folder)
        }
    except Exception as e:
        logger.error(f"[StorageInfo] Failed to get storage info for notebook {notebook_id}: {e}")
        return {"error": str(e)}


def _render_filters():
    return NotebookHelper.render_filters()


def _render_notebook_card(nb: store.Notebook, col):
    return NotebookHelper.render_notebook_card(nb, col)


def _note_title_from_content(content: str, max_len: int = 60) -> str:
    """Generate a concise human-friendly title from note content."""
    if not content:
        return "Untitled"
    try:
        text = content.strip()
        # Remove fenced code blocks
        text = re.sub(r"```[\s\S]*?```", " ", text)
        # Remove inline code backticks
        text = re.sub(r"`([^`]*)`", r"\1", text)
        # Replace links/images with their visible text
        text = re.sub(r"!\[([^\]]*)\]\([^\)]*\)", r"\1", text)  # images -> alt text
        text = re.sub(r"\[([^\]]+)\]\([^\)]*\)", r"\1", text)    # links -> label
        # Strip markdown headings prefixes
        text = re.sub(r"^\s*#+\s*", "", text, flags=re.MULTILINE)
        # First non-empty line
        first_line = next((ln.strip() for ln in text.splitlines() if ln.strip()), "")
        candidate = re.sub(r"\s+", " ", first_line) if first_line else re.sub(r"\s+", " ", text)
        candidate = candidate.strip()
        # Remove emphasis markers like **bold**, *italic*, __bold__, _italic_
        candidate = re.sub(r"\*\*([^*]+)\*\*", r"\1", candidate)
        candidate = re.sub(r"\*([^*]+)\*", r"\1", candidate)
        candidate = re.sub(r"__([^_]+)__", r"\1", candidate)
        candidate = re.sub(r"_([^_]+)_", r"\1", candidate)
        # Clean leftover multiple asterisks
        candidate = re.sub(r"\*{2,}", "", candidate)
        if len(candidate) > max_len:
            cutoff = candidate.rfind(" ", 0, max_len)
            if cutoff == -1:
                cutoff = max_len
            candidate = candidate[:cutoff].rstrip() + "..."
        return candidate or "Untitled"
    except Exception:
        return "Untitled"


def _render_sources_panel(nb: store.Notebook):
    """Render sources panel."""
    st.subheader("📚 Sources")

    # Ensure Studio notes are available for better note titles
    if 'studio_notes' not in st.session_state:
        st.session_state.studio_notes = {}
    if nb.id not in st.session_state.studio_notes:
        st.session_state.studio_notes[nb.id] = []
    if not st.session_state.studio_notes[nb.id]:
        try:
            loaded_notes = _load_notes_from_storage(nb.id)
            if loaded_notes:
                st.session_state.studio_notes[nb.id] = loaded_notes
        except Exception:
            pass
    
    # Lightweight cached helpers to avoid repeated IO on reruns
    @st.cache_data(show_spinner=False, ttl=3600)
    def _cached_page_title(url: str) -> str | None:
        try:
            return get_context()["document_processor"].get_page_title(url)
        except Exception:
            return None

    @st.cache_data(show_spinner=False, ttl=3600)
    def _cached_youtube_title(url: str) -> str | None:
        try:
            yinfo = get_context()["youtube_processor"].get_video_info(url)
            return yinfo.get("title")
        except Exception:
            return None

    @st.cache_data(show_spinner=False, ttl=300)
    def _cached_meta_search(source: str, notebook_id: str) -> list[dict]:
        try:
            ctx = get_context()
            return ctx['vector_db'].metadata_manager.search_metadata({
                'source': source,
                'notebook_id': notebook_id,
            }) if ctx.get('vector_db') else []
        except Exception:
            return []

    # Sources list
    if not nb.sources:
        st.info("No sources yet. Add files or links.")
    else:
        # Scrollable container for long source lists
        st.markdown('<div style="max-height:200px; overflow-y:auto; padding-right:8px;">', unsafe_allow_html=True)
        # Prepare metadata manager for chunk counts
        try:
            ctx = get_context()
            metadata_manager = ctx.get("vector_db").metadata_manager if ctx.get("vector_db") else None
        except Exception:
            metadata_manager = None

        for i, s in enumerate(nb.sources):
            col1, col2 = st.columns([4, 1])
            with col1:
                # Prefer Studio-like generated titles for note sources
                display_title = s.title
                display_href = None
                try:
                    if getattr(s, 'type', '') == 'note':
                        notes_list = st.session_state.studio_notes.get(nb.id, [])
                        matched = None
                        # Try by meta.note_id
                        meta = getattr(s, 'meta', None) or {}
                        note_id = meta.get('note_id') if isinstance(meta, dict) else None
                        if note_id:
                            for n in notes_list:
                                if n.get('id') == note_id or n.get('original_chat_id') == note_id:
                                    matched = n
                                    break
                        # Fallback by content prefix that was stored as source_path_or_url
                        if matched is None:
                            for n in notes_list:
                                prefix = (n.get('content') or "")[:100] + "..."
                                if getattr(s, 'source_path_or_url', None) == prefix:
                                    matched = n
                                    break
                        if matched is not None:
                            display_title = _note_title_from_content(matched.get('content', ''))
                        else:
                            # Fallback: try derive from stored snippet
                            display_title = _note_title_from_content(getattr(s, 'source_path_or_url', '') or s.title)
                    elif getattr(s, 'type', '') in ['url', 'youtube']:
                        # Attempt to fetch a human-friendly page title
                        url_value = getattr(s, 'source_path_or_url', s.title)
                        display_href = url_value
                        page_title = None
                        # Try using metadata stored in vector DB for this source (cached)
                        matches = _cached_meta_search(url_value, nb.id)
                        if matches:
                            meta0 = matches[0]
                            page_title = meta0.get('youtube_title') or meta0.get('page_title')
                        # Fallback: for youtube use processor to get title on-demand
                        if page_title is None and s.type == 'youtube':
                            page_title = _cached_youtube_title(url_value)
                        # Fallback: try fetch <title> from URL for generic web pages
                        if page_title is None and s.type == 'url':
                            page_title = _cached_page_title(url_value)
                        if page_title:
                            display_title = page_title
                        else:
                            display_title = url_value
                except Exception:
                    pass
                # Render as hyperlink if we have an href
                header_text = f"[{display_title}]({display_href})" if display_href else display_title
                with st.expander(f"{header_text}", expanded=False):
                    # Build rich details JSON for the source
                    chunk_count = 0
                    content_type = s.type
                    try:
                        if metadata_manager is not None:
                            src_key = getattr(s, 'source_path_or_url', s.title)
                            matches = _cached_meta_search(src_key, nb.id)
                            chunk_count = len(matches)
                            if matches:
                                content_type = matches[0].get('content_type', content_type)
                    except Exception:
                        pass
                    st.json({
                        'source': getattr(s, 'source_path_or_url', s.title),
                        'content_type': content_type,
                        'chunk_count': chunk_count,
                        'added_at': s.added_at,
                    })
                    if s.meta:
                        st.json(s.meta)
            with col2:
                if st.button("❌", key=f"del_source_{s.id}", help="Delete source"):
                    store.remove_source(nb.id, s.id)
                    st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)

    st.markdown("---")
    st.subheader("Add sources")
    uploaded = st.file_uploader(
        "Upload files",
        type=["mp4","avi","mov","mkv","mp3","wav","pdf","docx","txt","xlsx"],
        accept_multiple_files=True,
        key="nb_upload",
    )
    url = st.text_input(
        "Or add a link",
        key="nb_url",
        placeholder="Enter urls, youtube links,...",)
    # Internet search UI (keys are namespaced per-notebook to avoid cross-notebook leakage)
    st.markdown("Search Internet")
    with st.container():
        search_col, btn_col = st.columns([5,1])
        with search_col:
            search_query = st.text_input(
                "Search Internet",
                key=f"nb_web_search_{nb.id}",
                max_chars=100,
                placeholder="Enter keywords or topic (max 100 chars)",
                label_visibility="collapsed",
            )
        with btn_col:
            do_search = st.button("Search", key=f"nb_web_search_btn_{nb.id}", use_container_width=True)
    # Persist search results in session
    results_key = f"nb_web_results_{nb.id}"
    if results_key not in st.session_state:
        st.session_state[results_key] = []
    if do_search and search_query and search_query.strip():
        try:
            ctx = get_context()
            urls = ctx["web_search"].web_search(search_query.strip())
            st.session_state[results_key] = urls
            if not urls:
                st.info("No results found for your query.")
        except Exception as e:
            st.error(f"Web search failed: {e}")
            st.session_state[results_key] = []
    # Render selectable results
    if st.session_state.get(results_key):
        st.markdown("#### Search results")
        select_all = st.checkbox("Select all", key=f"nb_web_select_all_{nb.id}", value=False)
        selected_urls = []
        for idx, link in enumerate(st.session_state[results_key][:20]):
            checked = st.checkbox(link, key=f"nb_web_result_{nb.id}_{idx}", value=select_all)
            if checked:
                selected_urls.append(link)
    else:
        selected_urls = []
    if st.button("Add to notebook", key="nb_add_btn"):
        added = 0
        duplicates: list[str] = []
        # Build existing set of identifiers to avoid duplicates
        existing_keys = set()
        for s in nb.sources or []:
            key = (s.type, getattr(s, 'source_path_or_url', s.title))
            existing_keys.add(key)
        if uploaded:
            added += ingest_uploaded_files(nb.id, uploaded)
            for f in uploaded:
                key = ("file", f.name)
                if key in existing_keys:
                    duplicates.append(f.name)
                    continue
                store.add_source(nb.id, type="file", title=f.name, source_path_or_url=f.name)
                existing_keys.add(key)
        if url:
            src_type = ("youtube" if ("youtube.com" in url or "youtu.be" in url) else "url")
            key = (src_type, url)
            if key in existing_keys:
                duplicates.append(url)
            else:
                added += ingest_url(nb.id, url)
                # Try to resolve a better title for url/youtube
                resolved_title = url
                try:
                    if src_type == 'youtube':
                        yinfo = get_context()["youtube_processor"].get_video_info(url)
                        resolved_title = yinfo.get("title") or url
                    else:
                        page_title = get_context()["document_processor"].get_page_title(url)
                        resolved_title = page_title or url
                except Exception:
                    pass
                store.add_source(nb.id, type=src_type, title=resolved_title, source_path_or_url=url)
                existing_keys.add(key)
        # Add selected web search links
        for weblink in selected_urls:
            src_type = ("youtube" if ("youtube.com" in weblink or "youtu.be" in weblink) else "url")
            key = (src_type, weblink)
            if key in existing_keys:
                duplicates.append(weblink)
                continue
            try:
                resolved_title = weblink
                if src_type == 'youtube':
                    yinfo = get_context()["youtube_processor"].get_video_info(weblink)
                    resolved_title = yinfo.get("title") or weblink
                else:
                    page_title = get_context()["document_processor"].get_page_title(weblink)
                    resolved_title = page_title or weblink
                added += ingest_url(nb.id, weblink)
                store.add_source(nb.id, type=src_type, title=resolved_title, source_path_or_url=weblink)
                existing_keys.add(key)
            except Exception:
                continue
        if duplicates and added == 0:
            if len(duplicates) == 1:
                st.warning("Nguồn đã tồn tại trong notebook!")
            else:
                st.warning("Một số nguồn đã tồn tại trong notebook!")
        elif duplicates:
            st.warning("Một số nguồn đã tồn tại và đã được bỏ qua: " + ", ".join(duplicates[:3]) + ("..." if len(duplicates) > 3 else ""))
        # Clear web search results and query after any Add action to avoid repetition
        try:
            st.session_state[results_key] = []
            st.session_state[f"nb_web_search_{nb.id}"] = ""
            st.session_state[f"nb_web_select_all_{nb.id}"] = False
        except Exception:
            pass
        if added > 0:
            st.success(f"Added {added} chunks.")
            st.rerun()


def _notebook_overview(nb: store.Notebook):
    with st.expander("📘 Overview & Example questions", expanded=False):
        st.markdown("### Overview")

        # Check if we need to regenerate overview and examples
        current_sources_count = len(nb.sources)
        cached_sources_count = st.session_state.get(f"sources_count_{nb.id}", 0)
        stored_overview = store.get_overview(nb.id)
        stored_examples = nb.examples
        
        # Only regenerate if sources count changed or no cached data exists
        needs_regeneration = (
            current_sources_count != cached_sources_count or 
            not stored_overview or 
            not stored_examples
        )
        
        overview_task_key = f"studio_text_overview_task_{nb.id}"
        examples_task_key = f"studio_text_examples_task_{nb.id}"
        if overview_task_key not in _TASK_STATUS:
            _TASK_STATUS[overview_task_key] = {"running": False, "result": None, "error": None}
        if examples_task_key not in _TASK_STATUS:
            _TASK_STATUS[examples_task_key] = {"running": False, "result": None, "error": None}

        if needs_regeneration:
            # Trigger background generation once; UI remains responsive
            if not _TASK_STATUS[overview_task_key]["running"] and not _TASK_STATUS[examples_task_key]["running"]:
                _TASK_STATUS[overview_task_key] = {"running": True, "result": None, "error": None}
                _TASK_STATUS[examples_task_key] = {"running": True, "result": None, "error": None}
                threading.Thread(
                    target=_background_generate_overview_and_examples,
                    args=(nb.id, overview_task_key, examples_task_key),
                    daemon=True,
                ).start()
            # Show placeholders while background tasks run
            overview = stored_overview or "Đang tạo overview..."
            examples = stored_examples or ["Đang tạo câu hỏi ví dụ..."]
            # Update cached sources count so we don't retrigger until change
            st.session_state[f"sources_count_{nb.id}"] = current_sources_count
        else:
            # Use cached data
            overview = stored_overview
            examples = stored_examples
        
        st.write(overview)
        
        # Show cache status
        if not needs_regeneration:
            st.caption("📝 Overview và Examples đã được lưu cache")

        st.markdown("#### Example questions")
        cols = st.columns(len(examples))
        for i, ex in enumerate(examples[:3]):
            with cols[i]:
                if st.button(ex, key=f"ex_{i}"):
                    st.session_state["nb_chat_input"] = ex


def _generate_example_questions(nb: store.Notebook) -> List[str]:
    return NotebookHelper.generate_example_questions(nb)


def _render_chat(nb: store.Notebook):
    """Render chat interface with conversation view and save note per answer."""
    ctx = get_context()

    # Initialize chat histories per-notebook in session to avoid cross-notebook leakage
    if 'chat_histories' not in st.session_state:
        st.session_state.chat_histories = {}
    if nb.id not in st.session_state.chat_histories:
        st.session_state.chat_histories[nb.id] = []  # list of {id, question, answer, timestamp, sources}
    chat_history = st.session_state.chat_histories[nb.id]
    
    # Initialize button counter to ensure unique keys
    if 'chat_button_counter' not in st.session_state:
        st.session_state.chat_button_counter = 0

    # Preference for including sources (used when rendering history)
    include_sources_pref = st.session_state.get('nb_include_sources', True)

    # Ensure Studio notes store exists for this notebook
    if 'studio_notes' not in st.session_state:
        st.session_state.studio_notes = {}
    if nb.id not in st.session_state.studio_notes:
        st.session_state.studio_notes[nb.id] = []

    # Status placeholder
    status_ph = st.empty()

    # Show searching status when searching (avoid heavy re-render while toggling UI)
    if st.session_state.get('is_searching', False):
        status_ph.info("Searching and answering…")
    else:
        # Chat history display area - only show when not searching
        st.markdown("### Chat History")
        if chat_history:
            st.info(f"📝 Chat history has {len(chat_history)} messages")
            
            # Chat history styling
            from src.interface.utils.notebook_ui import NotebookUI
            chat_style = NotebookUI.chat_style_css(max_height=500)
            st.markdown(chat_style, unsafe_allow_html=True)
            st.markdown('<div class="nb-chat-wrapper">', unsafe_allow_html=True)
            
            for item in chat_history[-50:]:
                st.markdown(f'<div class="nb-chat-item nb-chat-q"><div class="bubble q-bubble">{item["question"]}</div></div>', unsafe_allow_html=True)
                st.markdown(f'<div class="nb-chat-item nb-chat-a"><div class="bubble a-bubble">{item["answer"]}</div></div>', unsafe_allow_html=True)
                
                # Action buttons for each message - now with 3 columns to accommodate Speak button
                bcols = st.columns([1,1,1,3])
                with bcols[0]:
                    # Stable key per chat message to avoid key drift across reruns
                    button_key = f"save_note_{item['id']}"
                    
                    # Check if this note was already saved
                    # Use per-notebook storage for studio notes
                    notes_for_this_nb = st.session_state.studio_notes.get(nb.id, [])
                    note_already_saved = any(n.get('original_chat_id') == item['id'] for n in notes_for_this_nb)
                    
                    if st.button("💾 Save Note" if not note_already_saved else "✅ Saved", 
                               key=button_key, 
                               disabled=note_already_saved):
                        # Ensure per-notebook notes store exists
                        if 'studio_notes' not in st.session_state:
                            st.session_state.studio_notes = {}
                        if nb.id not in st.session_state.studio_notes:
                            st.session_state.studio_notes[nb.id] = []
                        note = {
                            'id': f"note_{item['id']}",
                            'original_chat_id': item['id'],  # Track original chat message
                            'content': item['answer'],
                            'timestamp': datetime.now().isoformat(),
                            'sources': item.get('sources', []),
                            'question': item.get('question', ''),
                            'added_to_source': False
                        }
                        st.session_state.studio_notes[nb.id].append(note)
                        
                        # Save to persistent storage
                        _save_notes_to_storage(nb.id, st.session_state.studio_notes[nb.id])
                        
                        st.success("✅ Note saved to Studio!")
                        st.rerun()  # Rerun to update button state
                
                with bcols[1]:
                    # Speak button for text-to-speech
                    if st.button("🔊 Speak", key=f"speak_chat_{item['id']}", 
                               help="Listen to this answer", use_container_width=True):
                        try:
                            ctx = get_context()
                            if ctx.get("tts_client"):
                                with st.spinner("🎵 Generating audio..."):
                                    # Get answer content for TTS
                                    answer_text = item['answer']
                                    
                                    # Truncate text if too long (TTS has limits)
                                    max_length = 4000
                                    if len(answer_text) > max_length:
                                        answer_text = answer_text[:max_length] + "..."
                                        st.warning(f"⚠️ Answer was truncated for TTS (max {max_length} characters)")
                                    
                                    # Generate audio using TTS client
                                    audio_data = ctx["tts_client"].text_to_speech(
                                        text=answer_text,
                                        voice="alloy",  # Fixed voice
                                        model="tts-1",  # Fixed model - you can change this
                                        instructions="Speak clearly and at a moderate pace, suitable for answer reading."
                                    )
                                    
                                    if audio_data:
                                        # Create audio player with autoplay
                                        st.audio(audio_data, format="audio/mp3", start_time=0)
                                        st.success("🎵 Audio generated.")
                                    else:
                                        st.error("❌ Failed to generate audio")
                            else:
                                st.error("❌ TTS service not available")
                        except Exception as e:
                            st.error(f"❌ Error generating speech: {str(e)}")
                
                with bcols[2]:
                    if include_sources_pref and item.get('sources'):
                        with st.popover("Sources"):
                            for s in item['sources'][:5]:
                                st.markdown(f"- {s}")
            
            st.markdown('</div>', unsafe_allow_html=True)
        else:
            st.info("💬 No chat history yet. Ask a question to start chatting!")

    st.markdown("---")
    
    # Auto scroll to this section if flag is set
    if st.session_state.get('auto_scroll_to_question', False):
        st.session_state['auto_scroll_to_question'] = False
        # Add JavaScript for smooth scrolling
        from src.interface.utils.notebook_ui import NotebookUI
        st.markdown(NotebookUI.smooth_scroll_to_question_js(), unsafe_allow_html=True)
    
    st.markdown("### Ask a Question")

    # Group small controls together; put fast_mode first to minimize layout shifts
    col1, col2, col3 = st.columns([1,1,1])
    with col1:
        fast_mode = st.checkbox("Fast mode", value=False, help="Faster answers with smaller k and shorter generations")
    with col2:
        use_cot = st.checkbox("Chain-of-thought", value=False)
    with col3:
        include_src = st.checkbox("Include sources", value=include_sources_pref)
        st.session_state['nb_include_sources'] = include_src
        
    # Input + Ask in the same container and row
    with st.container(key="ask_question"):
        if st.session_state.get('clear_nb_input', False):
            st.session_state['nb_chat_input'] = ""
            st.session_state['clear_nb_input'] = False
        # Put a single label across both columns for better alignment
        input_col, ask_col = st.columns([9,1])
        with input_col:
            query = st.text_area("Your question", key="nb_chat_input", placeholder="Ask anything about the sources in this notebook…", label_visibility="collapsed")
        with ask_col:
            ask_clicked = st.button("Ask", type="primary", use_container_width=True)

    # Ask button handler
    if ask_clicked:
        if not query.strip():
            st.warning("Please enter a question")
            return
        
        # Set searching flag to hide chat history and show status
        st.session_state['is_searching'] = True
        
        try:
            # Smaller top-k in fast mode for speed
            k_value = 4 if fast_mode else 8
            results = ctx["search_engine"].search(query, k=k_value, threshold=0.2, filters={"notebook_id": nb.id})
            if not results:
                st.info("No relevant chunks found in this notebook.")
                # Reset searching flag to show chat history when no results
                st.session_state['is_searching'] = False
                return
            
            rag_results = [
                {"text": r.text or r.metadata.get("text", ""), "score": r.score, "metadata": r.metadata}
                for r in results
            ]
            # Detect question language and enforce consistent answer language
            vi_char_set = "àáạảãâầấậẩẫăằắặẳẵèéẹẻẽêềếệểễìíịỉĩòóọỏõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹđ"
            is_vietnamese = any(ch in vi_char_set for ch in query.lower())
            lang_instruction = (
                "Hãy trả lời bằng tiếng Việt, văn phong rõ ràng, mạch lạc."
                if is_vietnamese else
                "Answer in English with clear and concise language."
            )
            role_instruction = (
                "Bạn phải luôn trả lời bằng tiếng Việt."
                if is_vietnamese else
                "You must always respond in English."
            )
            # Prepare context: skip explicit summarization in fast mode to save time
            rag_for_sum = rag_results[:5] if fast_mode else rag_results
            if fast_mode:
                summary_text = "\n\n".join([r.get("text", "") for r in rag_for_sum])[:4000]
            else:
                # Use LangChain prompt + LLM to summarize when not in fast mode
                if ctx.get("prompt_manager") and ctx.get("llm_client"):
                    messages = ctx["prompt_manager"].build_summary_prompt(
                        content="\n\n".join([r.get("text", "") for r in rag_for_sum]),
                        additional_instructions=lang_instruction,
                    )
                    summary_resp = ctx["llm_client"].generate_response(
                        messages,
                        max_tokens=600,
                    )
                    summary_text = summary_resp.get("content", "")
                elif ctx.get("legacy_summarizer"):
                    summary_result = ctx["legacy_summarizer"].summarize_chunks(
                        rag_for_sum,
                        query=query,
                        use_chain_of_thought=use_cot,
                        extra_instructions=lang_instruction,
                        role=role_instruction,
                    )
                    summary_text = summary_result.summary

            # Use LangChain as primary AI system
            from src.utils.feature_flags import feature_flags

            answer = summary_text
            structured_output = None
            try:
                # Build prompt using LangChain prompt manager
                messages = ctx["prompt_manager"].build_qa_prompt(
                    context=summary_text,
                    question=query,
                    use_chain_of_thought=use_cot and not fast_mode,
                )
                messages.insert(0, {"role": "system", "content": role_instruction})

                # Use LangChain LLM client
                response = ctx["llm_client"].generate_response(
                    messages,
                    use_memory=not fast_mode,
                    store_in_memory=not fast_mode,
                    max_memory_context=3,
                    context_sources=[r.metadata.get("source","unknown") for r in results],
                    # Per-call overrides for latency in fast mode
                    max_tokens=600 if fast_mode else None,
                )
                # Normalize response format
                answer = response.get("content") if isinstance(response, dict) else getattr(response, "content", answer)
                
                # Parse structured output if available
                if feature_flags.is_enabled("use_structured_output") and ctx.get("output_parser"):
                    try:
                        structured_output = ctx["output_parser"].parse_qa(answer)
                        sources = [r.metadata.get("source", "unknown") for r in results[:5]]
                        structured_output.citations = ctx["output_parser"].extract_citations(answer, sources)
                    except Exception:
                        pass
            except Exception as _e:
                # Fallback to legacy system if LangChain fails
                try:
                    if ctx.get("legacy_prompt_engineer") and ctx.get("legacy_llm_client"):
                        pe = ctx["legacy_prompt_engineer"]
                        messages = pe.build_qa_prompt(
                            context=summary_text,
                            question=query,
                            use_chain_of_thought=use_cot and not fast_mode,
                        )
                        messages.insert(0, {"role": "system", "content": role_instruction})
                        
                        response = ctx["legacy_llm_client"].generate_response(
                            messages,
                            use_memory=not fast_mode,
                            store_in_memory=not fast_mode,
                            max_memory_context=3,
                            context_sources=[r.metadata.get("source","unknown") for r in results],
                        )
                        answer = response.content
                except Exception:
                    # Keep summary_text as final fallback
                    pass

            # Append new message to chat history with structured output
            new_message = {
                'id': f"chat_{datetime.now().strftime('%Y%m%d_%H%M%S_%f')}",
                'question': query,
                'answer': answer,
                'timestamp': datetime.now().isoformat(),
                'sources': [r.metadata.get('source','unknown') for r in results[:5]],
                'structured_output': structured_output
            }
            # Append to this notebook's chat history only
            st.session_state.chat_histories[nb.id].append(new_message)
            
            # Show success message
            st.success("✅ Answer generated! Scroll up to see the chat history.")
            
            # Reset searching flag to show chat history
            st.session_state['is_searching'] = False
            
            # Set flags for UI behavior on rerun
            st.session_state['auto_scroll_to_question'] = True
            st.session_state['clear_nb_input'] = True
            
            # Force rerun to show updated chat history
            st.rerun()
            
        except Exception as e:
            st.error(f"Error generating answer: {str(e)}")
            # Reset searching flag to show chat history after error
            st.session_state['is_searching'] = False


def _render_studio_actions(nb: store.Notebook, layout: str | None = None):
    """Render action buttons.
    layout: 'docx_only' | 'audio_only' | None (both rows previously).
    """
    task_docx_key = f"studio_docx_task_{nb.id}"
    task_audio_key = f"studio_audio_task_{nb.id}"

    # Initialize task states
    if task_docx_key not in _TASK_STATUS:
        _TASK_STATUS[task_docx_key] = {"running": False, "file_path": None, "error": None}
    if task_audio_key not in _TASK_STATUS:
        _TASK_STATUS[task_audio_key] = {"running": False, "file_path": None, "error": None}

    if layout in (None, "docx_only"):
        # DOCX button (not full width)
        docx_state = _TASK_STATUS.get(task_docx_key, {"running": False})
        docx_label = "⏳ Generating DOCX..." if docx_state.get("running") else "📄 Tổng quan bằng file DOCX"
        if st.button(docx_label, key=f"btn_docx_{nb.id}", disabled=docx_state.get("running", False)):
            _TASK_STATUS[task_docx_key] = {"running": True, "file_path": None, "error": None}
            threading.Thread(target=_background_generate_docx_overview, args=(nb.id, task_docx_key), daemon=True).start()
        if docx_state.get("error"):
            st.error(f"❌ Lỗi: {docx_state['error']}")

    if layout in (None, "audio_only"):
        # AUDIO button (not full width)
        audio_state = _TASK_STATUS.get(task_audio_key, {"running": False})
        audio_label = "⏳ Generating audio..." if audio_state.get("running") else "🔊 Tổng quan bằng âm thanh"
        if st.button(audio_label, key=f"btn_audio_{nb.id}", disabled=audio_state.get("running", False)):
            _TASK_STATUS[task_audio_key] = {"running": True, "file_path": None, "error": None}
            threading.Thread(target=_background_generate_audio_overview, args=(nb.id, task_audio_key), daemon=True).start()
        if audio_state.get("error"):
            st.error(f"❌ Lỗi: {audio_state['error']}")


def _render_studio_panel(nb: store.Notebook):
    """Render studio panel with saved notes (no file manager UI)."""

    # Initialize per-notebook Studio notes store
    if 'studio_notes' not in st.session_state:
        st.session_state.studio_notes = {}
    if nb.id not in st.session_state.studio_notes:
        st.session_state.studio_notes[nb.id] = []
    
    # Load notes from persistent storage if session state is empty
    if not st.session_state.studio_notes[nb.id]:
        stored_notes = _load_notes_from_storage(nb.id)
        if stored_notes:
            st.session_state.studio_notes[nb.id] = stored_notes
            st.success(f"📚 Loaded {len(stored_notes)} notes from storage")

    # One-time migration from legacy saved_notes list (if present)
    if 'saved_notes' in st.session_state and st.session_state.saved_notes:
        st.session_state.studio_notes[nb.id].extend(st.session_state.saved_notes)
        st.session_state.saved_notes = []
        # Save migrated notes to storage
        _save_notes_to_storage(nb.id, st.session_state.studio_notes[nb.id])

    notes = st.session_state.studio_notes[nb.id]
    # Deduplicate notes by original_chat_id (fallback id)
    try:
        unique_notes_map = {}
        for n in notes:
            dedup_key = n.get('original_chat_id') or n.get('id')
            if dedup_key not in unique_notes_map:
                unique_notes_map[dedup_key] = n
        if len(unique_notes_map) != len(notes):
            st.session_state.studio_notes[nb.id] = list(unique_notes_map.values())
            notes = st.session_state.studio_notes[nb.id]
            # Save deduplicated notes to storage
            _save_notes_to_storage(nb.id, notes)
    except Exception:
        pass

    # Sync Studio notes with current sources: if a previously added note's source was deleted,
    # re-enable the Add to Source button by resetting flags.
    try:
        existing_source_ids = {s.id for s in (nb.sources or [])}
        changed = False
        for n in notes:
            if not n.get('added_to_source'):
                continue
            src_id = n.get('source_id')
            # If we have a stored source_id but it no longer exists, try to relink or reset
            if src_id and src_id not in existing_source_ids:
                matched = None
                # Try to find by meta.note_id if available
                note_identifier = n.get('id') or n.get('original_chat_id')
                for s in (nb.sources or []):
                    meta = getattr(s, 'meta', None) or {}
                    if isinstance(meta, dict) and meta.get('note_id') == note_identifier:
                        matched = s
                        break
                # Fallback: match by content prefix used when creating source
                if matched is None:
                    prefix = (n.get('content') or "")[:100] + "..."
                    for s in (nb.sources or []):
                        try:
                            if s.type == 'note' and s.source_path_or_url == prefix:
                                matched = s
                                break
                        except Exception:
                            continue
                if matched is not None:
                    n['source_id'] = matched.id
                    n['added_to_source'] = True
                    changed = True
                else:
                    n['added_to_source'] = False
                    n.pop('source_id', None)
                    changed = True
            # If we do not have a stored source_id but note is marked added, try to verify existence
            elif not src_id:
                matched = None
                note_identifier = n.get('id') or n.get('original_chat_id')
                for s in (nb.sources or []):
                    meta = getattr(s, 'meta', None) or {}
                    if isinstance(meta, dict) and meta.get('note_id') == note_identifier:
                        matched = s
                        break
                if matched is None:
                    prefix = (n.get('content') or "")[:100] + "..."
                    for s in (nb.sources or []):
                        try:
                            if s.type == 'note' and s.source_path_or_url == prefix:
                                matched = s
                                break
                        except Exception:
                            continue
                if matched is not None:
                    n['source_id'] = matched.id
                    n['added_to_source'] = True
                    changed = True
                else:
                    n['added_to_source'] = False
                    changed = True
        if changed:
            _save_notes_to_storage(nb.id, notes)
    except Exception:
        pass

    if not notes:
        st.info("No saved notes yet. Save notes from chat to see them here.")
        return
    
    # Display saved notes
    for i, note in enumerate(notes):
        generated_name = _note_title_from_content(note.get('content', ''))
        with st.expander(f"📝 {generated_name} (Note {i+1} - {note['timestamp'][:16]})", expanded=False):
            st.markdown(note['content'])
            st.caption(f"Sources: {', '.join(note['sources'][:2])}")
            
            # Action buttons - now with 3 columns to accommodate Speak button
            col1, col2, col3 = st.columns([2, 1, 1])
            
            with col1:
                add_label = "📚 Add to Source" if not note.get('added_to_source') else "✅ Added"
                if st.button(add_label, key=f"add_to_source_{note['id']}_{i}", 
                           help="Convert note to source", disabled=note.get('added_to_source', False)):
                    # Add note as a new source
                    # Compose a better source title from content, fallback to date
                    generated_name = _note_title_from_content(note.get('content', ''))
                    title_text = generated_name or f"Note from {note['timestamp'][:10]}"
                    sref = store.add_source(
                        nb.id, 
                        type="note", 
                        title=title_text, 
                        source_path_or_url=note['content'][:100] + "...",
                        meta={
                            "origin": "studio_note",
                            "note_id": note.get('id') or note.get('original_chat_id')
                        }
                    )
                    # Mark note as added to sources but keep it in Studio
                    note['added_to_source'] = True
                    if sref is not None:
                        note['source_id'] = sref.id
                    # Save updated notes to storage
                    _save_notes_to_storage(nb.id, st.session_state.studio_notes[nb.id])
                    st.success("✅ Note added to sources! (kept in Studio)")
                    st.rerun()
            
            with col2:
                # Speak button for text-to-speech
                if st.button("🔊 Speak", key=f"speak_note_{note['id']}_{i}", 
                           help="Listen to this note", use_container_width=True):
                    try:
                        # Get context and TTS client with fallback
                        ctx = get_context()
                        tts_client = ctx.get("tts_client")
                        
                        # If TTS client is None, try to re-initialize it
                        if tts_client is None:
                            st.info("🔄 Re-initializing TTS client...")
                            try:
                                from src.ai.tts_client import TTSClient
                                tts_client = TTSClient()
                                # Update context
                                ctx["tts_client"] = tts_client
                                st.success("✅ TTS client re-initialized!")
                            except Exception as init_error:
                                st.error(f"❌ Failed to re-initialize TTS client: {init_error}")
                                return
                        
                        if tts_client:
                            with st.spinner("🎵 Generating audio..."):
                                # Get note content for TTS
                                note_text = note['content']
                                    
                                # Truncate text if too long (TTS has limits)
                                max_length = 4000
                                if len(note_text) > max_length:
                                    note_text = note_text[:max_length] + "..."
                                    st.warning(f"⚠️ Note was truncated for TTS (max {max_length} characters)")
                                    
                                # Generate audio using TTS client with fixed model
                                audio_data = tts_client.text_to_speech(
                                    text=note_text,
                                    voice="alloy",  # Fixed voice
                                    model="tts-1",  # Fixed model - you can change this
                                    instructions="Speak clearly and at a moderate pace, suitable for note reading."
                                )
                                
                                if audio_data:
                                    # Create audio player with autoplay
                                    st.audio(audio_data, format="audio/mp3", start_time=0)
                                    st.success("🎵 Audio generated")
                                else:
                                    st.error("❌ Failed to generate audio")
                        else:
                            st.error("❌ TTS service not available after re-initialization")
                    except Exception as e:
                        st.error(f"❌ Error generating speech: {str(e)}")
                        st.info("💡 This might be due to API rate limits or network issues")
            
            with col3:
                if st.button("🗑️ Delete", key=f"delete_note_{note['id']}_{i}", 
                           help="Delete note"):
                    st.session_state.studio_notes[nb.id].pop(i)
                    # Save updated notes to storage
                    _save_notes_to_storage(nb.id, st.session_state.studio_notes[nb.id])
                    st.success("✅ Note deleted!")
                    st.rerun()


def _render_create_view():
    edit_notebook_id = st.session_state.get("edit_notebook_id")
    is_edit_mode = edit_notebook_id is not None

    if is_edit_mode:
        st.title("Edit notebook")
        notebook = store.get_notebook(edit_notebook_id)
        if not notebook:
            st.error("Notebook not found")
            st.query_params["view"] = "list"
            st.rerun()
            return
    else:
        st.title("Create a new notebook")

    with st.form("create_notebook_form"):
        name = st.text_input(
            "Notebook name",
            placeholder="e.g., Marketing Q4 Reports",
            value=notebook.name if is_edit_mode else ""
        )
        desc = st.text_area(
            "Description (optional)",
            value=notebook.description if is_edit_mode else ""
        )
        tags = st.text_input(
            "Tags (comma-separated)",
            value=", ".join(notebook.tags) if is_edit_mode and notebook.tags else ""
        )
        uploaded = st.file_uploader(
            "Upload files",
            type=["mp4","avi","mov","mkv","mp3","wav","pdf","docx","txt","xlsx"],
            accept_multiple_files=True,
        )
        url = st.text_input("Add a link (web page or YouTube)")
        submitted = st.form_submit_button("Save" if is_edit_mode else "Create", type="primary")

    if submitted:
        if not name.strip():
            st.error("Please provide a notebook name")
            return

        if is_edit_mode:
            store.update_notebook(
                edit_notebook_id,
                name=name.strip(),
                description=desc.strip(),
                tags=[t.strip() for t in tags.split(',') if t.strip()]
            )
            st.success("Notebook updated successfully!")
            st.session_state.pop("edit_notebook_id", None)
            st.session_state["current_notebook_id"] = edit_notebook_id
            st.query_params["view"] = "notebook"
            st.rerun()
        else:
            nb = store.create_notebook(name=name.strip(), description=desc.strip(), tags=[t.strip() for t in tags.split(',') if t.strip()])
            added = 0
            if uploaded:
                added += ingest_uploaded_files(nb.id, uploaded)
                for f in uploaded:
                    store.add_source(nb.id, type="file", title=f.name, source_path_or_url=f.name)
            if url:
                added += ingest_url(nb.id, url)
                store.add_source(nb.id, type=("youtube" if ("youtube.com" in url or "youtu.be" in url) else "url"), title=url, source_path_or_url=url)

            st.success(f"Notebook created. Added {added} chunks to knowledge base.")
            st.session_state["current_notebook_id"] = nb.id
            st.query_params["view"] = "notebook"
            st.rerun()

    if st.button("← Back to Notebooks"):
        st.session_state.pop("edit_notebook_id", None)
        st.query_params["view"] = "list"
        st.rerun()


def main():
    st.set_page_config(page_title="Notebooks", page_icon="📓", layout="wide")
    
    # Load custom CSS
    with open("src/interface/styles/notebook_layout.css") as f:
        st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)

    action = st.query_params.get("action", None)
    notebook_id = st.query_params.get("id", None)
    view = st.query_params.get("view", "list")

    if action == "open" and notebook_id:
        st.session_state["current_notebook_id"] = notebook_id
        st.query_params["view"] = "notebook"
        st.rerun()
    elif action == "favorite" and notebook_id:
        store.toggle_favorite(notebook_id)
        st.rerun()
    elif action == "edit" and notebook_id:
        st.session_state["edit_notebook_id"] = notebook_id
        st.query_params["view"] = "create"
        st.rerun()
    elif action == "delete" and notebook_id:
        st.session_state[f"confirm_del_{notebook_id}"] = True

    if view == "notebook":
        nb_id = st.session_state.get("current_notebook_id")
        if not nb_id:
            st.warning("No notebook selected. Returning to list.")
            st.query_params["view"] = "list"
            st.rerun()
            return
        nb = store.get_notebook(nb_id)
        if not nb:
            st.warning("Notebook not found. Returning to list.")
            st.query_params["view"] = "list"
            st.rerun()
            return

        # Header
        header_cols = st.columns([9,1])
        with header_cols[0]:
            st.title(nb.name)
            st.caption(nb.description or "")
            # Info chips: sources count • tags • created date
            created_date = (nb.created_at or "").split("T")[0]
            sources_count = len(nb.sources) if nb.sources else 0
            tags_text = ", ".join(nb.tags[:5]) if getattr(nb, 'tags', None) else "No tags"
            st.markdown(
                f"**📄 {sources_count} sources** • **🏷️ {tags_text}** • **📅 {created_date}**"
            )
        with header_cols[1]:
            if st.button("📓 Notebooks"):
                st.query_params["view"] = "list"
                st.rerun()

        # Three-tab layout: Notebook | Studio | Source (Settings moved under Studio)
        tab_notebook, tab_studio, tab_sources = st.tabs(["📓 **Notebook**", "🎨 **Studio**", "📚 **Source**"])

        with tab_notebook:
            _notebook_overview(nb)
            st.divider()
            _render_chat(nb)

        with tab_studio:
            with st.expander("**🎥 Studio**", expanded=True):
                # Compact 3x2 grid: Generate | Download (adjacent)
                # Row 1: DOCX
                r1c1, r1c2 = st.columns([10, 1])
                with r1c1:
                    _render_studio_actions(nb, layout="docx_only")
                with r1c2:
                    latest_docx = _get_latest_docx_path(nb.id)
                    if latest_docx.exists():
                        with open(latest_docx, "rb") as f:
                            st.download_button("⬇️", f, file_name=latest_docx.name, help="Tải báo cáo DOCX")

                # Row 2: Audio
                r2c1, r2c2 = st.columns([10, 1])
                with r2c1:
                    _render_studio_actions(nb, layout="audio_only")
                with r2c2:
                    latest_audio = _get_latest_audio_path(nb.id)
                    if latest_audio.exists():
                        with open(latest_audio, "rb") as f:
                            st.download_button("⬇️", f, file_name=latest_audio.name, help="Tải audio")

                # Row 3: Mindmap (placeholder)
                r3c1, r3c2 = st.columns([10, 1])
                with r3c1:
                    st.button("🧠 Bản đồ tư duy (TODO)", disabled=True, key=f"btn_mindmap_{nb.id}")
                with r3c2:
                    latest_map = _get_latest_mindmap_path(nb.id)
                    if latest_map.exists():
                        with open(latest_map, "rb") as f:
                            st.download_button("⬇️", f, file_name=latest_map.name, help="Tải mindmap")

            # Notes section (renamed from Studio inner content)
            with st.expander("**📝 Notes**", expanded=True):
                _render_studio_panel(nb)
            st.divider()
            with st.expander("**⚙️ Settings**", expanded=False):
                with st.form("notebook_settings_form"):
                    new_name = st.text_input("Đổi tên notebook", value=nb.name)
                    new_tags_str = st.text_input("Edit tag của notebook", value=", ".join(nb.tags) if getattr(nb, 'tags', None) else "")
                    submitted_settings = st.form_submit_button("Lưu cài đặt", type="primary")
                if submitted_settings:
                    if not new_name.strip():
                        st.error("Vui lòng nhập tên notebook hợp lệ")
                    else:
                        new_tags_list = [t.strip() for t in new_tags_str.split(',') if t.strip()]
                        store.update_notebook(
                            nb.id,
                            name=new_name.strip(),
                            description=nb.description or "",
                            tags=new_tags_list,
                        )
                        st.success("Đã cập nhật cài đặt notebook")
                        st.rerun()

        with tab_sources:
            _render_sources_panel(nb)

        return

    if view == "create":
        _render_create_view()
        return

    st.title("📓 ElevateAI Notebooks")
    st.markdown("Create, organize and chat with your knowledge notebooks.")

    create_col = st.container()
    with create_col:
        if st.button("Create New Notebook", type="primary"):
            st.query_params["view"] = "create"
            st.rerun()

    st.markdown("---")
    st.subheader("Your Notebooks")

    # Lazy load filters and notebooks
    with st.spinner("Loading notebooks..."):
        q, fav, dfrom, dto = _render_filters()
        
        # Cache notebooks in session state to avoid reloading
        cache_key = f"notebooks_cache_{hash(f'{q}_{fav}_{dfrom}_{dto}')}"
        if cache_key not in st.session_state:
            st.session_state[cache_key] = store.list_notebooks(q, fav, dfrom, dto)
        
        notebooks = st.session_state[cache_key]
        
        if not notebooks:
            st.info("No notebooks yet. Create your first notebook to get started!")
            return
        
        # Lazy render notebooks with pagination
        page_size = 8  # Show 8 notebooks per page
        current_page = st.session_state.get('notebooks_page', 0)
        total_pages = (len(notebooks) + page_size - 1) // page_size
        
        # Pagination controls
        if total_pages > 1:
            col1, col2, col3 = st.columns([1, 2, 1])
            with col1:
                if st.button("← Previous", disabled=current_page == 0):
                    st.session_state['notebooks_page'] = max(0, current_page - 1)
                    st.rerun()
            with col2:
                st.write(f"Page {current_page + 1} of {total_pages}")
            with col3:
                if st.button("Next →", disabled=current_page >= total_pages - 1):
                    st.session_state['notebooks_page'] = min(total_pages - 1, current_page + 1)
                    st.rerun()
        
        # Calculate start and end indices for current page
        start_idx = current_page * page_size
        end_idx = min(start_idx + page_size, len(notebooks))
        current_notebooks = notebooks[start_idx:end_idx]
        
        # Render notebooks in grid with lazy loading
        cols = st.columns(4)
        for i, nb in enumerate(current_notebooks):
            col_index = i % 4
            with cols[col_index]:
                # Use container to isolate each notebook card
                with st.container():
                    _render_notebook_card(nb, cols[col_index])
        
        # Show total count
        st.caption(f"Showing {start_idx + 1}-{end_idx} of {len(notebooks)} notebooks")
        
        # Clear cache button for debugging (disable during searching)
        if st.button("🔄 Refresh Notebooks", help="Clear cache and reload notebooks", disabled=st.session_state.get('is_searching', False)):
            # Clear all notebook caches
            for key in list(st.session_state.keys()):
                if key.startswith('notebooks_cache_'):
                    del st.session_state[key]
            st.session_state['notebooks_page'] = 0
            st.rerun()


if __name__ == "__main__":
    main()


